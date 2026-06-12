from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "manuscript"
FIGURES = MANUSCRIPT / "figures"
PACKAGE = ROOT / "submission_package"
PACKAGE_FIGURES = PACKAGE / "figures"

TITLE = "Operationalizing the waste hierarchy for climate-ready organic waste transitions"


def _font(size: int, bold: bool = False) -> ImageFont.ImageFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


def _add_panel(canvas: Image.Image, path: Path, box: tuple[int, int, int, int], label: str) -> None:
    img = Image.open(path).convert("RGB")
    x, y, w, h = box
    img.thumbnail((w, h), Image.Resampling.LANCZOS)
    px = x + (w - img.width) // 2
    py = y + (h - img.height) // 2 + 18
    canvas.paste(img, (px, py))
    draw = ImageDraw.Draw(canvas)
    draw.text((x + 10, y + 8), label, fill=(25, 25, 25), font=_font(42, bold=True))


def compose_figures() -> None:
    PACKAGE_FIGURES.mkdir(parents=True, exist_ok=True)

    specs = {
        "Figure_1.png": [
            ("a", "map_ofmsw_generation.png", (60, 40, 1640, 880)),
            ("b", "map_first_pass_methane.png", (1740, 40, 1640, 880)),
            ("c", "fig_region_pathway_totals.png", (260, 1000, 2860, 860)),
        ],
        "Figure_2.png": [
            ("a", "fig_global_pathway_totals.png", (80, 40, 1120, 640)),
            ("b", "map_best_pathway.png", (1260, 40, 2080, 900)),
            ("c", "fig_best_pathway_counts.png", (80, 760, 1120, 520)),
            ("d", "fig_top20_country_hotspots.png", (420, 1340, 2600, 720)),
        ],
        "Figure_3.png": [
            ("a", "fig_uncertainty_intervals.png", (80, 40, 1500, 720)),
            ("b", "fig_win_probability_by_winner.png", (1740, 40, 1500, 720)),
            ("c", "fig_robust_winner_counts.png", (120, 920, 1100, 620)),
            ("d", "fig_robust_recovery_opportunities.png", (1420, 880, 1720, 820)),
        ],
        "Figure_4.png": [
            ("a", "map_readiness_typology.png", (80, 40, 1980, 900)),
            ("b", "fig_readiness_scatter.png", (2140, 80, 1240, 860)),
            ("c", "fig_top_readiness_countries.png", (80, 1080, 1640, 900)),
            ("d", "fig_readiness_weighting_robustness.png", (1880, 1180, 1360, 700)),
        ],
    }

    for filename, panels in specs.items():
        canvas = Image.new("RGB", (3500, 2200), "white")
        for label, source, box in panels:
            _add_panel(canvas, FIGURES / source, box, label)
        canvas.save(PACKAGE_FIGURES / filename, quality=95)


def clean_inline(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = text.replace("`", "")
    return text


def add_markdown_paragraph(doc: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        return
    if stripped.startswith("|"):
        p = doc.add_paragraph(clean_inline(stripped))
        p.style = "Normal"
        return
    if stripped.startswith("- "):
        p = doc.add_paragraph(clean_inline(stripped[2:]), style="List Bullet")
        return
    if re.match(r"^\d+\.\s", stripped):
        p = doc.add_paragraph(clean_inline(re.sub(r"^\d+\.\s", "", stripped)), style="List Number")
        return
    if stripped.startswith("### "):
        doc.add_heading(clean_inline(stripped[4:]), level=3)
    elif stripped.startswith("## "):
        doc.add_heading(clean_inline(stripped[3:]), level=2)
    elif stripped.startswith("# "):
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(clean_inline(stripped[2:]))
        run.bold = True
        run.font.size = Pt(16)
    else:
        p = doc.add_paragraph(clean_inline(stripped))
        p.style = "Normal"


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Arial"


def markdown_to_docx(markdown: str, out: Path, include_figures: bool = False) -> None:
    doc = Document()
    configure_doc(doc)

    lines = markdown.splitlines()
    for line in lines:
        add_markdown_paragraph(doc, line)

    if include_figures:
        doc.add_page_break()
        doc.add_heading("Embedded figures", level=1)
        for idx in range(1, 5):
            path = PACKAGE_FIGURES / f"Figure_{idx}.png"
            doc.add_paragraph(f"Figure {idx}")
            doc.add_picture(str(path), width=Inches(6.5))
            if idx < 4:
                doc.add_page_break()

    doc.save(out)


def build_main_markdown() -> str:
    article = (MANUSCRIPT / "nature_sustainability_article_draft.md").read_text()
    references = (MANUSCRIPT / "references_draft.md").read_text()
    references = references.split("## Literature Still To Add")[0].replace("# References Draft", "## References")
    return f"{article.rstrip()}\n\n{references.strip()}\n"


def build_readme() -> None:
    text = """# Nature Sustainability upload package

Use these files for the submission system:

1. `NatureSustainability_Analysis_Main_Manuscript_with_Figures.docx`
   - File type: Manuscript
   - Contains main text, references, figure captions, and draft embedded figures.

2. `NatureSustainability_Analysis_Supplementary_Information.docx`
   - File type: Supplementary Information

3. `NatureSustainability_Analysis_Cover_Letter.docx`
   - File type: Cover Letter

The figure composites in `figures/` were regenerated from the updated ecoinvent 3.12 pathway outputs.

Do not upload raw ecoinvent archives, ecoSpold files, or licensed screenshots/tables as source data.
"""
    (PACKAGE / "UPLOAD_README.md").write_text(text)


def main() -> None:
    PACKAGE.mkdir(exist_ok=True)
    compose_figures()
    markdown_to_docx(
        build_main_markdown(),
        PACKAGE / "NatureSustainability_Analysis_Main_Manuscript_with_Figures.docx",
        include_figures=True,
    )
    markdown_to_docx(
        (MANUSCRIPT / "supplementary_information_draft.md").read_text(),
        PACKAGE / "NatureSustainability_Analysis_Supplementary_Information.docx",
    )
    markdown_to_docx(
        (MANUSCRIPT / "nature_sustainability_cover_letter_draft.md").read_text(),
        PACKAGE / "NatureSustainability_Analysis_Cover_Letter.docx",
    )
    build_readme()


if __name__ == "__main__":
    main()
